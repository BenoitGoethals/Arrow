"""COP document management — upload, list, download, and delete supporting files."""

from __future__ import annotations

from datetime import datetime, timezone

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy.orm import Session

from backend.audit import log_event
from backend.auth.jwt_auth import get_current_operator, require_role
from backend.missions.dependencies import get_active_mission
from backend.storage.database import get_db
from backend.storage.models import CopDocument, Mission, Operator

router = APIRouter(prefix="/cop", tags=["cop"])

ALLOWED_EXTENSIONS = {".txt", ".doc", ".docx"}
MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MB


class CopDocumentOut(BaseModel):
    id: int
    title: str
    filename: str
    content_type: str
    text_preview: str
    uploaded_by: int
    uploaded_at: datetime
    mission_id: int | None

    model_config = {"from_attributes": True}


@router.get("/documents", response_model=list[CopDocumentOut])
def list_documents(
    db: Session = Depends(get_db),
    _: Operator = Depends(get_current_operator),
    mission: Mission | None = Depends(get_active_mission),
) -> list[CopDocument]:
    q = db.query(CopDocument).order_by(CopDocument.uploaded_at.desc())
    if mission:
        q = q.filter(CopDocument.mission_id == mission.id)
    return q.limit(200).all()


@router.post("/documents", response_model=CopDocumentOut, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current: Operator = Depends(get_current_operator),
    mission: Mission | None = Depends(get_active_mission),
) -> CopDocument:
    filename = file.filename or "upload"
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type '{ext}' not allowed. Accepted: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    data = await file.read()
    if len(data) > MAX_FILE_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds 10 MB limit")

    content_type = "TXT" if ext == ".txt" else "WORD"
    text_preview = ""
    if content_type == "TXT":
        text_preview = data.decode("utf-8", errors="replace")[:20_000]
    else:
        # Best-effort Word text extraction; requires python-docx if installed.
        try:
            import io
            from docx import Document as DocxDocument  # type: ignore[import]
            doc = DocxDocument(io.BytesIO(data))
            text_preview = "\n".join(p.text for p in doc.paragraphs)[:20_000]
        except Exception:
            text_preview = ""

    title = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").title()
    doc = CopDocument(
        title=title,
        filename=filename,
        content_type=content_type,
        file_data=data,
        text_preview=text_preview,
        uploaded_by=current.id,
        mission_id=mission.id if mission else current.mission_id,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    log_event(db, "COP_DOCUMENT_UPLOADED", operator_id=current.id,
              resource=f"cop_doc:{doc.id}", detail=f"file:{filename}")
    return doc


@router.get("/documents/{doc_id}/download")
def download_document(
    doc_id: int,
    db: Session = Depends(get_db),
    _: Operator = Depends(get_current_operator),
) -> Response:
    doc = db.get(CopDocument, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")
    mime = "text/plain" if doc.content_type == "TXT" else (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    return Response(
        content=doc.file_data,
        media_type=mime,
        headers={"Content-Disposition": f'attachment; filename="{doc.filename}"'},
    )


@router.delete("/documents/{doc_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_document(
    doc_id: int,
    db: Session = Depends(get_db),
    current: Operator = Depends(require_role("ADMIN", "BATTLE_CAPTAIN")),
) -> None:
    doc = db.get(CopDocument, doc_id)
    if doc is None:
        raise HTTPException(status_code=404, detail="Document not found")
    db.delete(doc)
    db.commit()
    log_event(db, "COP_DOCUMENT_DELETED", operator_id=current.id, resource=f"cop_doc:{doc_id}")
